import re, html, os, json, shutil, datetime
import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__))) if "__file__" in dir() else "."
SRC_DIR = os.path.join(REPO_ROOT, "source-docs")
OUT_DIR = REPO_ROOT

DOCS = [
    {"docx": "2001-16.docx", "title": "Unit Titles Act 2001", "slug": "unit-titles-act-2001",
     "citation": "A2001-16", "kind": "Act", "abbrev": "UTA"},
    {"docx": "2001-58.docx", "title": "Community Title Act 2001", "slug": "community-title-act-2001",
     "citation": "A2001-58", "kind": "Act", "abbrev": "CTA"},
    {"docx": "2011-39.docx", "title": "Unit Titles (Management) Regulation 2011", "slug": "unit-titles-management-regulation-2011",
     "citation": "SL2011-39", "kind": "Regulation", "abbrev": "UTMR"},
    {"docx": "2011-41.docx", "title": "Unit Titles (Management) Act 2011", "slug": "unit-titles-management-act-2011",
     "citation": "A2011-41", "kind": "Act", "abbrev": "UTMA"},
]

# ACT Parliamentary Counsel's Word template uses named paragraph styles that
# directly encode document structure — no running headers/footers, no
# line-wrapped headings, no sentence/heading ambiguity to guess at. This maps
# every style seen across all 4 documents to how it should be handled.

# style name -> (heading level, id-kind). Levels are the real document
# hierarchy, so a section always nests one level below the Part/Division/
# Subdivision (or Schedule/Schedule Part) that contains it.
STRUCT_STYLES = {
    "A H2 Part": (2, "part"),
    "A H3 Div": (3, "div"),
    "A H4 SubDiv": (4, "subdiv"),
    "Sched-heading": (2, "sch"),
    "Sched-Part": (3, "schpart"),
    "Dict-Heading": (2, "fixed:dictionary"),
    "Endnote1": (2, "fixed:endnotes"),
    "Endnote2": (3, "endnote"),
}

SECTION_STYLES = {"A H5 Sec", "Sch clause heading"}

CSS_CLASS_BY_STYLE = {
    "A main": "lead", "Sch A main": "lead", "LongTitle": "lead",
    "A main return": "lead lead-return",
    "A para": "clause clause-a", "Sch A para": "clause clause-a",
    "A subpara": "clause clause-i", "Sch A subpara": "clause clause-i",
    "A subsubpara": "clause clause-A",
    "aDef": "def", "aDef para": "def", "aDef subpara": "def",
    "ref": "ref",
    "Penalty": "penalty",
    "New Act": "amdt-act", "New Reg": "amdt-act",
    "Act details": "amdt-detail", "As am by": "amdt-detail",
    "AmdtsEntryHd": "amdt-head",
    "AmdtsEntries": "amdt-entry", "AmdtsEntriesDefL2": "amdt-entry",
    "EndNoteTextPub": "", "EndNoteTextEPS": "",
    "Normal": "",
    "Formula": "",
}
# Any style starting with these prefixes gets the given class.
CSS_CLASS_PREFIX = {
    "aNote": "note",
    "aExam": "example",
}

# Styles that never contribute body content (front matter, table of
# contents, empty structural bookmarks).
SKIP_STYLES = {
    "Billname", "Billname1", "BillBasic", "ActNo", "RepubNo", "EffectiveDate",
    "CoverInForce", "CoverHeading", "CoverSubHdg", "CoverText", "CoverTextBullet",
    "CoverActName", "N-TOCheading", "N-9pt", "PageBreak", "Placeholder",
    "00SigningPage", "01Contents", "02Text", "03Schedule", "04Dictionary",
    "05EndNote", "06Copyright", "N-line3",
}

def is_skippable(style_name, text):
    if style_name.startswith("toc"):
        return True
    if style_name in SKIP_STYLES:
        return True
    if style_name == "Normal" and (text == "Australian Capital Territory" or text.startswith("©")):
        return True
    return False

def slugify_num(num):
    return re.sub(r'[.\s]+', '-', num)

def split_num_title(text):
    """Heading paragraphs are stored as '<number>\\t<title>'."""
    if "\t" in text:
        num, title = text.split("\t", 1)
        return num.strip(), title.strip()
    parts = text.split(None, 1)
    if len(parts) == 2:
        return parts[0], parts[1]
    return text, ""

NUM_TAIL_RE = re.compile(r'(\d+(?:\.\d+)*[A-Za-z]*)$')

def bare_num(label):
    """Part/Division/Subdivision/Schedule headings store the label word in
    front of the number ('Division 3.2') — strip it for id-building."""
    m = NUM_TAIL_RE.search(label)
    return m.group(1) if m else label

def extract_metadata(doc):
    meta = {}
    for p in doc.paragraphs:
        t = p.text.strip()
        if p.style.name == "RepubNo":
            m = re.search(r'\d+', t)
            if m:
                meta["republication_no"] = m.group(0)
        elif p.style.name == "EffectiveDate":
            meta["effective"] = t.split(":", 1)[-1].strip().replace("\xa0", " ")
        elif p.style.name == "CoverInForce" and t.lower().startswith("last amendment made by"):
            meta["last_amendment"] = t.split("by", 1)[-1].strip()
    return meta

def iter_body(doc):
    """Paragraphs and tables in document order. Schedules of reviewable
    decisions, the endnote abbreviation key and the earlier-republications
    list are all tables, so walking doc.paragraphs alone drops them."""
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, doc)
        elif child.tag == qn('w:tbl'):
            yield Table(child, doc)

def is_header_style(style_name):
    return style_name.endswith("ColHd") or style_name.endswith("Hdg")

def read_table(table):
    """Rows of cells, each cell a list of paragraph strings, plus a flag for
    whether the row is a column-heading row."""
    rows = []
    for row in table.rows:
        cells, seen = [], set()
        for cell in row.cells:
            key = id(cell._tc)
            if key in seen:      # horizontally merged cell, already emitted
                continue
            seen.add(key)
            lines = []
            for p in cell.paragraphs:
                t = p.text.strip().replace("\t", " ")
                if t:
                    lines.append(t)
            cells.append(lines)
        if not any(cells):
            continue
        styles = [p.style.name for c in row.cells for p in c.paragraphs if p.text.strip()]
        rows.append({"cells": cells, "header": bool(styles) and all(is_header_style(s) for s in styles)})
    return rows

def parse_doc(doc, unmapped_styles):
    blocks = []
    schedule_scope = ""   # e.g. "sch-1-"
    subpart_scope = ""    # e.g. "part-1-1-" (nested inside a schedule)
    scope = []            # [(level, "Part 3"), (level, "Division 3.3")]
    seen_struct_ids = {}

    def unique_id(hid):
        n = seen_struct_ids.get(hid, 0) + 1
        seen_struct_ids[hid] = n
        return hid if n == 1 else f"{hid}-{n}"

    for item in iter_body(doc):
        if isinstance(item, Table):
            rows = read_table(item)
            if rows:
                blocks.append({"type": "table", "kind": "table", "rows": rows})
            continue

        p = item
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name
        if is_skippable(style, text):
            continue

        if style in STRUCT_STYLES:
            level, kind = STRUCT_STYLES[style]
            label = None
            if kind == "part":
                schedule_scope = subpart_scope = ""
                num, title = split_num_title(text)
                hid = f"part-{slugify_num(bare_num(num))}"
                label = num
            elif kind == "div":
                num, title = split_num_title(text)
                hid = f"div-{slugify_num(bare_num(num))}"
                label = num
            elif kind == "subdiv":
                num, title = split_num_title(text)
                hid = f"subdiv-{slugify_num(bare_num(num))}"
                label = num
            elif kind == "sch":
                num, title = split_num_title(text)
                subpart_scope = ""
                hid = f"sch-{slugify_num(bare_num(num))}"
                schedule_scope = hid + "-"
                label = num
            elif kind == "schpart":
                num, title = split_num_title(text)
                part_num = slugify_num(bare_num(num))
                hid = schedule_scope + f"part-{part_num}"
                subpart_scope = f"part-{part_num}-"
                label = num
            elif kind == "fixed:dictionary":
                title, hid = "Dictionary", "dictionary"
                schedule_scope = subpart_scope = ""
            elif kind == "fixed:endnotes":
                title, hid = "Endnotes", "endnotes"
                schedule_scope = subpart_scope = ""
            elif kind == "endnote":
                num, t = split_num_title(text)
                title = f"{num} {t}"
                hid = f"endnote-{slugify_num(num)}"

            while scope and scope[-1][0] >= level:
                scope.pop()
            if label:
                scope.append((level, label))

            display = text.replace("\t", " ") if kind not in ("fixed:dictionary", "fixed:endnotes") else title
            blocks.append({"type": f"h{level}", "kind": "struct", "text": display,
                           "id": unique_id(hid), "level": level})
            continue

        if style in SECTION_STYLES:
            num, title = split_num_title(text)
            hid = schedule_scope + subpart_scope + f"s-{slugify_num(num)}"
            # Schedule clause numbering restarts in each Part (UTMA sch 1 has a
            # clause 8 in both pt 1.1 and pt 1.2), and a bare "26 ..." heading
            # doesn't say which Part or Division it sits under. The breadcrumb
            # keeps every heading self-identifying when it is read out of context.
            level = min(5, (scope[-1][0] + 1) if scope else 3)
            blocks.append({"type": f"h{level}", "kind": "section", "text": f"{num} {title}",
                           "id": unique_id(hid), "level": level,
                           "scope": ", ".join(lbl for _, lbl in scope)})
            continue

        cls = CSS_CLASS_BY_STYLE.get(style)
        if cls is None:
            for prefix, c in CSS_CLASS_PREFIX.items():
                if style.startswith(prefix):
                    cls = c
                    break
        if cls is None:
            unmapped_styles.add(style)
            cls = ""
        blocks.append({"type": "p", "kind": "content", "text": text.replace("\t", " "), "class": cls})

    return blocks

def render_toc(nav_items):
    """Nested list of the structural skeleton (Parts, Divisions, Subdivisions,
    Schedules, Schedule Parts, Dictionary, Endnotes). Sections are deliberately
    left out: repeating 250-odd section titles here would compete with the
    provisions themselves when the page is indexed."""
    out, depth, open_li = ["<ul>"], 1, False
    for hid, text, level in nav_items:
        while depth < level - 1:
            out.append("<ul>")       # nested list belongs inside the open <li>
            depth += 1
            open_li = False
        while depth > level - 1:
            if open_li:
                out.append("</li>")
                open_li = False
            out.append("</ul>")
            depth -= 1
            open_li = True           # the <li> that wrapped the nested list
        if open_li:
            out.append("</li>")
        out.append(f'<li><a href="#{hid}">{text}</a>')
        open_li = True
    while depth > 0:
        if open_li:
            out.append("</li>")
            open_li = False
        out.append("</ul>")
        depth -= 1
        open_li = depth > 0
    return "\n".join(out)

def render_table(rows):
    out = ['<div class="table-wrap">', "<table>"]
    body_open = False
    for i, row in enumerate(rows):
        if row["header"]:
            if i == 0:
                out.append("<thead>")
            tag = "th"
        else:
            if i > 0 and rows[0]["header"] and not body_open:
                out.append("</thead>")
                out.append("<tbody>")
                body_open = True
            tag = "td"
        cells = "".join(
            f"<{tag}>" + "<br>".join(html.escape(line) for line in cell) + f"</{tag}>"
            for cell in row["cells"]
        )
        out.append(f"<tr>{cells}</tr>")
    if body_open:
        out.append("</tbody>")
    out.append("</table>")
    out.append("</div>")
    return "\n".join(out)

# --- Output layout ---------------------------------------------------------
#
# The site is a hub with the legislation reader nested under it, not a bare set of documents:
#
#   /                     the hub, generated from the strata-kit manifest
#   /legislation/         the reader's index
#   /legislation/*.html   one page per instrument
#   /assets/*.css         one stylesheet per page type, kit theme concatenated in
#   /<slug>.html          redirect stubs, because the documents used to live at the root
#
# Absolute asset paths are safe here because the site is served at a custom domain's root.
# They would break on the bare github.io repository subpath, which this site no longer uses.

KIT_DIR = os.path.join(REPO_ROOT, "vendor", "strata-kit")
LEG_DIR = os.path.join(REPO_ROOT, "legislation")
ASSETS_DIR = os.path.join(REPO_ROOT, "assets")
SITE_URL = "https://strata.noradz.io"
BUILD_DATE = datetime.date.today().isoformat()


def load_kit():
    """The family manifest and the pre-rendered navigation bar, from the strata-kit submodule.

    The bar is rendered once in the kit and committed there rather than re-implemented here:
    the other consumers are Vite builds, and two renderers for one piece of markup drift.
    """
    with open(os.path.join(KIT_DIR, "projects.json")) as f:
        manifest = json.load(f)
    with open(os.path.join(KIT_DIR, "nav.html")) as f:
        nav = f.read().strip()
    return manifest, nav


def kit_nav(nav, current):
    """Mark this site's own entry in the shared bar. The kit ships one file for every site."""
    marker = f'data-kit-id="{current}"'
    return nav.replace(marker, f'{marker} aria-current="page"', 1)


# The glyph is noradz's own nav mark (a circle and its chord — see the noradz site's own
# favicon), kept in the family's rose-on-white treatment to match the sibling sites' icons.
FAVICON = (
    "data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 32 32'%3E"
    "%3Crect width='32' height='32' rx='8' fill='%23b33a5f'/%3E"
    "%3Ccircle cx='16' cy='17' r='6' fill='none' stroke='%23fff' stroke-width='2'/%3E"
    "%3Cpath d='M11.13 13.5L20.87 13.5' stroke='%23fff' stroke-width='2' "
    "stroke-linecap='round'/%3E%3C/svg%3E"
)


def page(title, description, stylesheet, nav, body, extra_head=""):
    """The shell every page in this repo shares: head, the family bar, then the page's body."""
    return f"""<!doctype html>
<html lang="en-AU">
<head>
<meta charset="utf-8">
<title>{title}</title>
<meta name="description" content="{description}">
<meta name="viewport" content="width=device-width, initial-scale=1">
<link rel="icon" href="{FAVICON}">
<link rel="stylesheet" href="/assets/{stylesheet}">
{extra_head}</head>
<body>
<a class="skip-link" href="#main">Skip to content</a>
{nav}
{body}
</body>
</html>
"""


THEME_TOGGLE = """<button class="theme-toggle" type="button" data-theme-toggle aria-pressed="false">
<span class="visually-hidden">Switch between the light and dark palette</span>
<svg width="15" height="15" viewBox="0 0 24 24" fill="none" aria-hidden="true">
<circle cx="12" cy="12" r="4.2" stroke="currentColor" stroke-width="1.8"></circle>
<path d="M12 2.4v2.2M12 19.4v2.2M2.4 12h2.2M19.4 12h2.2M5.2 5.2l1.6 1.6M17.2 17.2l1.6 1.6M18.8 5.2l-1.6 1.6M6.8 17.2l-1.6 1.6" stroke="currentColor" stroke-width="1.8" stroke-linecap="round"></path>
</svg>
<span data-theme-label>Light</span>
</button>"""


def render_html(doc_meta, blocks, page_meta, nav):
    nav_items = []
    body_html = []
    for b in blocks:
        if b["kind"] == "table":
            body_html.append(render_table(b["rows"]))
            continue
        esc = html.escape(b["text"])
        if b["kind"] in ("struct", "section"):
            tag = b["type"]
            scope = b.get("scope")
            extra = f' <span class="heading-scope">({html.escape(scope)})</span>' if scope else ""
            cls = ' class="struct"' if b["kind"] == "struct" else ""
            body_html.append(f'<{tag}{cls} id="{b["id"]}">{esc}{extra}</{tag}>')
            if b["kind"] == "struct":
                nav_items.append((b["id"], esc, b["level"]))
        else:
            cls = f' class="{b["class"]}"' if b["class"] else ""
            body_html.append(f'<p{cls}>{esc}</p>')

    toc_html = render_toc(nav_items)

    meta_bits = [f'{doc_meta["citation"]} ({doc_meta["kind"]})',
                 f'cited here as {doc_meta["abbrev"]}',
                 "Australian Capital Territory"]
    if page_meta.get("effective"):
        meta_bits.append(f'Effective: {html.escape(page_meta["effective"])}')
    if page_meta.get("republication_no"):
        meta_bits.append(f'Republication No {html.escape(page_meta["republication_no"])}')
    if page_meta.get("last_amendment"):
        meta_bits.append(f'Last amendment: {html.escape(page_meta["last_amendment"])}')
    meta_line = " &middot; ".join(meta_bits)

    other_docs = "\n".join(
        f'<li><a href="{d["slug"]}.html">{html.escape(d["title"])} ({d["abbrev"]})</a></li>'
        for d in DOCS if d["slug"] != doc_meta["slug"]
    )

    body = f"""<nav class="doc-switch" aria-label="Other documents">
<ul>
{other_docs}
</ul>
</nav>
<main id="main">
<h1>{html.escape(doc_meta["title"])} ({doc_meta["abbrev"]})</h1>
<p class="doc-meta">{meta_line}</p>
<p class="source-note">This page reproduces the full text of the <a href="https://www.legislation.act.gov.au/">ACT legislation register</a> republication. It is not an authorised legal copy &mdash; always confirm current provisions at <a href="https://www.legislation.act.gov.au/">legislation.act.gov.au</a>.</p>
<nav class="toc" aria-label="Table of contents">
<h2 class="toc-heading">Contents</h2>
{toc_html}
</nav>
<article>
{chr(10).join(body_html)}
</article>
</main>
<footer class="site-footer">
<p><a href="index.html">&larr; All four documents</a> &middot; <a href="/">strata.noradz.io</a></p>
</footer>"""

    return page(
        title=f'{html.escape(doc_meta["title"])} ({doc_meta["abbrev"]}) — ACT strata legislation',
        description=f'Full text of the {html.escape(doc_meta["title"])} ({doc_meta["citation"]}), Australian Capital Territory, cited as {doc_meta["abbrev"]}.',
        stylesheet="legislation.css",
        nav=nav,
        body=body,
    )


def build_legislation_index(nav):
    items = "\n".join(
        f'<li><a href="{d["slug"]}.html">{html.escape(d["title"])}</a> '
        f'<span class="citation">({d["citation"]}, cited as {d["abbrev"]})</span></li>'
        for d in DOCS
    )
    body = f"""<main id="main">
<h1>ACT strata legislation</h1>
<p>The full text of the Australian Capital Territory legislation governing unit titles, community title, and owners corporation management. Each document below is reproduced in full on its own page, with the document's own heading hierarchy intact so a provision can be linked to directly.</p>
<ul class="doc-list">
{items}
</ul>
<p class="source-note">Source: <a href="https://www.legislation.act.gov.au/">ACT legislation register</a>. These are unofficial reproductions kept for reference; always confirm current provisions at the official register.</p>
</main>
<footer class="site-footer">
<p><a href="/">&larr; strata.noradz.io</a></p>
</footer>"""
    return page(
        title="ACT strata legislation — full text",
        description="The full text of the ACT unit titles and community title legislation, one page per instrument.",
        stylesheet="legislation.css",
        nav=nav,
        body=body,
    )


def build_hub(manifest, nav):
    """The landing page: what the family is, and a card for each project in the manifest.

    The cards are the manifest — adding a strata project means adding an entry in the kit and
    moving this repo's submodule pin, not editing markup here.
    """
    family = manifest["family"]
    cards = []
    for project in manifest["projects"]:
        if not project.get("hub"):
            continue
        status = project["status"]
        chip = "" if status == "live" else f'<span class="chip chip-{html.escape(status)}">{html.escape(status)}</span>'
        heading = html.escape(project["name"])
        link_open, link_close = (f'<a href="{html.escape(project["url"])}">', "</a>") if status != "planned" else ("", "")
        visit = "" if status == "planned" else f"""<a class="card-visit" href="{html.escape(project["url"])}">
Open
<svg width="13" height="13" viewBox="0 0 24 24" fill="none" aria-hidden="true"><path d="M7 17 17 7M9 7h8v8" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"/></svg>
</a>"""
        cards.append(f"""<li class="card card-{html.escape(project["id"])}">
<h2>{link_open}{heading}{link_close}{chip}</h2>
<p class="card-summary">{html.escape(project["summary"])}</p>
<p class="card-blurb">{html.escape(project["blurb"])}</p>
{visit}
</li>""")

    body = f"""<div class="shell">
<header class="masthead">
<h1 class="wordmark">{html.escape(family["name"])}<span class="dot">.</span></h1>
<p class="tagline">{html.escape(family["tagline"])}</p>
{THEME_TOGGLE}
</header>
<main id="main">
<ul class="cards">
{chr(10).join(cards)}
</ul>
<section class="about">
<h2>What this is</h2>
<p>A side project, not a product.</p>
<p>The tools run entirely in your browser. Nothing you open in them is uploaded anywhere, there is no account, and there is no analytics — each of them ships a content security policy that stops the page making a network request at all, which you can check in your browser's network tab.</p>
<p class="caveat">None of it is legal advice, and the reproduced legislation is unofficial. For anything that matters, confirm the current provision at the <a href="https://www.legislation.act.gov.au/">ACT legislation register</a>.</p>
</section>
</main>
</div>
<script type="module" src="/assets/theme-toggle.js"></script>"""

    return page(
        title="strata — ACT strata legislation and tools",
        description=family["tagline"],
        stylesheet="hub.css",
        nav=nav,
        body=body,
    )


def build_redirect(slug):
    """A stub at each document's old root-level URL.

    The pages moved under /legislation/ when the hub took the root. The stub redirects in
    script first, because that is the only way to carry the fragment across — the deep links
    that matter here are anchors to a single section (#s-26), and a meta refresh drops them.
    The meta refresh is the fallback for a client with no script, and the visible link is the
    fallback for one that honours neither.
    """
    target = f"/legislation/{slug}.html"
    return f"""<!doctype html>
<html lang="en-AU">
<head>
<meta charset="utf-8">
<title>Moved — {slug}</title>
<link rel="canonical" href="{SITE_URL}{target}">
<meta name="robots" content="noindex">
<meta http-equiv="refresh" content="0; url={target}">
<script>location.replace("{target}" + location.hash);</script>
</head>
<body>
<p>This page has moved to <a href="{target}">{target}</a>.</p>
</body>
</html>
"""


def build_sitemap(manifest):
    urls = [f"{SITE_URL}/", f"{SITE_URL}/legislation/"]
    urls += [f"{SITE_URL}/legislation/{d['slug']}.html" for d in DOCS]
    entries = "\n".join(
        f"  <url>\n    <loc>{u}</loc>\n    <lastmod>{BUILD_DATE}</lastmod>\n  </url>" for u in urls
    )
    return f"""<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
{entries}
</urlset>
"""


# --- Styles ----------------------------------------------------------------
#
# Each page type gets one stylesheet, built by concatenating the family fonts, the kit's
# theme.css and the page's own rules. Concatenating rather than linking three files keeps it to
# one request and removes any question of which order they load in; the kit theme is a few
# kilobytes, so carrying it in both files costs nothing worth optimising.

FONTS_CSS = """/* Self-hosted, not from a CDN: the family's sites make no third-party requests, and this
   one holds to that even though it has no CSP of its own to enforce it. Both faces are
   OFL-licensed; the licences are alongside the files in assets/fonts/. */
@font-face {
  font-family: 'Fraunces Variable';
  font-style: normal;
  font-display: swap;
  font-weight: 100 900;
  src: url('fonts/fraunces-latin-full-normal.woff2') format('woff2-variations');
  unicode-range: U+0000-00FF, U+0131, U+0152-0153, U+02BB-02BC, U+02C6, U+02DA, U+02DC, U+2000-206F, U+2074, U+20AC, U+2122, U+2191, U+2193, U+2212, U+2215, U+FEFF, U+FFFD;
}
@font-face {
  font-family: 'DM Sans';
  font-style: normal;
  font-display: swap;
  font-weight: 400;
  src: url('fonts/dm-sans-latin-400-normal.woff2') format('woff2');
  unicode-range: U+0000-00FF, U+0131, U+0152-0153, U+02BB-02BC, U+02C6, U+02DA, U+02DC, U+2000-206F, U+2074, U+20AC, U+2122, U+2191, U+2193, U+2212, U+2215, U+FEFF, U+FFFD;
}
@font-face {
  font-family: 'DM Sans';
  font-style: normal;
  font-display: swap;
  font-weight: 500;
  src: url('fonts/dm-sans-latin-500-normal.woff2') format('woff2');
  unicode-range: U+0000-00FF, U+0131, U+0152-0153, U+02BB-02BC, U+02C6, U+02DA, U+02DC, U+2000-206F, U+2074, U+20AC, U+2122, U+2191, U+2193, U+2212, U+2215, U+FEFF, U+FFFD;
}

"""

BASE_CSS = """
* { box-sizing: border-box; }

body {
  background: var(--ground);
  color: var(--ink);
  font-family: var(--font-body);
  line-height: 1.55;
  margin: 0;
  padding: 0;
  -webkit-font-smoothing: antialiased;
}

.skip-link {
  position: absolute;
  left: -999px;
  top: 0;
  background: var(--surface-sunk);
  color: var(--ink);
  padding: 0.5rem 1rem;
}
.skip-link:focus { left: 0; z-index: 30; }

a { color: var(--accent-text); }

"""

# The reader's own vocabulary, expressed in the kit's tokens rather than in colours of its own.
# Keeping the local names means the typographic rules below are untouched by the change of
# palette, and a future token rename is one block to edit rather than eighty declarations.
LEGISLATION_CSS = """
:root {
  --muted: var(--ink-faint);
  --accent-bg: var(--surface-sunk);
  --struct: var(--accent-text);
}

main {
  max-width: 46rem;
  margin: 0 auto;
  padding: 1.5rem 1.25rem 4rem;
}
h1, h2, h3, h4, h5 { line-height: 1.25; font-family: var(--font-display); font-variation-settings: 'SOFT' 30; letter-spacing: -0.01em; }
h1 { font-size: 1.9rem; margin-top: 0.5rem; }
h2 { font-size: 1.35rem; margin-top: 2.4rem; border-top: 1px solid var(--border); padding-top: 1.2rem; }
h3 { font-size: 1.15rem; margin-top: 1.8rem; }
h4 { font-size: 1.05rem; margin-top: 1.5rem; }
h5 { font-size: 1rem; margin-top: 1.3rem; }
/* Part/Division/Subdivision/Schedule headings are dividers, not provisions —
   set them apart so a section never reads as a sibling of its own container. */
h3.struct, h4.struct, h5.struct {
  font-family: var(--font-body);
  color: var(--struct);
  text-transform: uppercase;
  letter-spacing: 0.04em;
  font-size: 0.85rem;
  margin-top: 2rem;
  padding-bottom: 0.3rem;
  border-bottom: 1px solid var(--border);
}
.heading-scope { font-weight: 400; font-size: 0.8rem; color: var(--muted); }
p { margin: 0.8rem 0; }
p.lead { margin-top: 1rem; }
p.lead-return { margin-top: 0.4rem; }
p.clause-a { margin-left: 1.5rem; }
p.clause-i { margin-left: 3rem; }
p.clause-A { margin-left: 4.5rem; }
p.def { margin-left: 1.5rem; font-style: italic; }
p.note, p.example { margin-left: 1.5rem; color: var(--muted); font-size: 0.92rem; font-style: italic; }
p.penalty { margin-left: 1.5rem; font-size: 0.95rem; }
p.ref { color: var(--muted); font-size: 0.85rem; margin: 0.2rem 0 0.8rem; }
p.amdt-act { font-weight: 600; margin-top: 1rem; }
p.amdt-head { font-weight: 600; margin-top: 0.8rem; }
p.amdt-detail, p.amdt-entry { color: var(--muted); font-size: 0.88rem; margin: 0.15rem 0 0.15rem 1.5rem; }
p.source-note, p.doc-meta { color: var(--muted); font-size: 0.92rem; }
.table-wrap { overflow-x: auto; margin: 1rem 0; }
table { border-collapse: collapse; width: 100%; font-size: 0.92rem; }
th, td { border: 1px solid var(--border); padding: 0.4rem 0.6rem; text-align: left; vertical-align: top; }
th { background: var(--accent-bg); font-weight: 600; }
.doc-switch {
  max-width: 46rem;
  margin: 0.75rem auto 0;
  padding: 0 1.25rem;
}
.doc-switch ul { display: flex; flex-wrap: wrap; gap: 0.75rem 1.25rem; list-style: none; padding: 0; margin: 0; font-size: 0.9rem; }
.toc {
  background: var(--surface);
  border: 1px solid var(--border);
  border-radius: var(--radius-card);
  padding: 1rem 1.25rem;
  margin: 1.5rem 0;
}
.toc-heading { margin: 0 0 0.5rem; border: none; padding: 0; font-size: 1rem; }
.toc ul { margin: 0; padding-left: 1.1rem; }
.toc ul ul { padding-left: 1.2rem; font-size: 0.94rem; }
.toc li { margin: 0.15rem 0; }
.doc-list { list-style: none; padding: 0; }
.doc-list li { padding: 0.6rem 0; border-bottom: 1px solid var(--border); }
.citation { color: var(--muted); font-size: 0.9rem; }
.site-footer { max-width: 46rem; margin: 0 auto; padding: 1rem 1.25rem 3rem; font-size: 0.9rem; }
"""

HUB_CSS = """
/* The wash at the top of the page, carried over from lodger so the hub reads as the same
   product as the tools it links to. */
body::before {
  content: '';
  position: fixed;
  inset: 0 0 auto;
  height: 420px;
  background:
    radial-gradient(60% 100% at 12% 0%, color-mix(in srgb, var(--rose) 46%, transparent), transparent 70%),
    radial-gradient(52% 100% at 88% 0%, color-mix(in srgb, var(--bleuet) 42%, transparent), transparent 70%),
    radial-gradient(40% 80% at 50% 0%, color-mix(in srgb, var(--lemon) 34%, transparent), transparent 75%);
  pointer-events: none;
  z-index: 0;
}

.shell {
  position: relative;
  z-index: 1;
  max-width: 900px;
  margin: 0 auto;
  padding: 40px 24px 96px;
}

.masthead {
  position: relative;
  margin-bottom: 36px;
  padding-right: 120px;
}

.wordmark {
  margin: 0;
  font-family: var(--font-display);
  font-optical-sizing: auto;
  font-variation-settings: 'SOFT' 40, 'WONK' 1;
  font-weight: 600;
  font-size: clamp(2.6rem, 6vw, 3.6rem);
  letter-spacing: -0.02em;
  line-height: 1;
}

.wordmark .dot { color: var(--accent-text); }

.tagline {
  margin: 10px 0 0;
  max-width: 54ch;
  color: var(--ink-soft);
  font-size: 1.05rem;
}

.theme-toggle {
  position: absolute;
  top: 4px;
  right: 0;
  display: inline-flex;
  align-items: center;
  gap: 7px;
  padding: 7px 13px;
  border: 1px solid var(--border);
  border-radius: 999px;
  background: var(--surface);
  color: var(--ink-soft);
  font-family: var(--font-body);
  font-size: 0.8rem;
  cursor: pointer;
  box-shadow: var(--shadow-soft);
}
.theme-toggle:hover { color: var(--ink); border-color: var(--border-strong); }
.theme-toggle:focus-visible { outline: 2px solid var(--accent); outline-offset: 3px; }

.cards {
  list-style: none;
  margin: 0;
  padding: 0;
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(260px, 1fr));
  gap: 20px;
}

.card {
  position: relative;
  background: var(--surface);
  border: 1px solid var(--border);
  border-radius: var(--radius-card);
  box-shadow: var(--shadow-soft);
  padding: 22px 22px 60px;
  /* The macaron edge is what distinguishes one card from the next at a glance; it is a
     surface, never a text colour, which is the rule that keeps the pastel palette readable. */
  border-top: 4px solid var(--border-strong);
}
.card-legislation { border-top-color: var(--lavender); }
.card-lodger { border-top-color: var(--rose); }
.card-former { border-top-color: var(--bleuet); }

.card h2 {
  margin: 0 0 8px;
  font-family: var(--font-display);
  font-variation-settings: 'SOFT' 30;
  font-weight: 600;
  font-size: 1.3rem;
  letter-spacing: -0.01em;
}
.card h2 a { text-decoration: none; color: var(--ink); }
.card h2 a:hover { color: var(--accent-text); }

.chip {
  margin-left: 8px;
  padding: 2px 9px;
  border-radius: 999px;
  background: var(--lemon);
  color: var(--ink);
  font-family: var(--font-body);
  font-size: 0.7rem;
  font-weight: 500;
  text-transform: lowercase;
  vertical-align: middle;
}

.card-summary { margin: 0 0 10px; color: var(--ink); font-weight: 500; }
.card-blurb { margin: 0; color: var(--ink-soft); font-size: 0.92rem; }

.card-visit {
  /* Absolute against the tile itself, not the text flow — so it lands in the same corner of
     every card regardless of blurb length, and doesn't depend on flexbox stretch behaviour
     that some browsers (Safari included) get wrong on a grid item. */
  position: absolute;
  right: 22px;
  bottom: 22px;
  display: inline-flex;
  align-items: center;
  gap: 5px;
  padding: 6px 13px;
  border-radius: 999px;
  /* Filled with the same macaron as the card's top edge — a surface, never a text colour,
     so it pairs with --ink rather than carrying the colour itself. */
  color: var(--ink);
  font-family: var(--font-body);
  font-size: 0.78rem;
  font-weight: 500;
  text-decoration: none;
  box-shadow: var(--shadow-soft);
  transition: box-shadow 0.15s ease, transform 0.15s ease;
}
.card-legislation .card-visit { background: var(--lavender); }
.card-lodger .card-visit { background: var(--rose); }
.card-former .card-visit { background: var(--bleuet); }
.card-visit svg { flex: none; }
.card-visit:hover { box-shadow: var(--shadow-lift); transform: translate(1px, -1px); }
.card-visit:focus-visible { outline: 2px solid var(--accent); outline-offset: 3px; }

.about {
  margin-top: 48px;
  max-width: 62ch;
}
.about h2 {
  font-family: var(--font-display);
  font-variation-settings: 'SOFT' 30;
  font-weight: 600;
  font-size: 1.15rem;
  margin: 0 0 10px;
}
.about p { color: var(--ink-soft); margin: 0 0 12px; }
.about .caveat { color: var(--ink-faint); font-size: 0.9rem; }

@media (prefers-reduced-motion: reduce) {
  * { animation-duration: 0.01ms !important; transition-duration: 0.01ms !important; }
}
"""


def main():
    manifest, nav_template = load_kit()
    with open(os.path.join(KIT_DIR, "theme.css")) as f:
        theme_css = f.read()

    os.makedirs(LEG_DIR, exist_ok=True)
    os.makedirs(ASSETS_DIR, exist_ok=True)

    leg_nav = kit_nav(nav_template, "legislation")
    hub_nav = nav_template  # The hub is the family's home; the bar's wordmark already points here.

    unmapped_styles = set()
    for doc_meta in DOCS:
        path = os.path.join(SRC_DIR, doc_meta["docx"])
        doc = docx.Document(path)
        page_meta = extract_metadata(doc)
        blocks = parse_doc(doc, unmapped_styles)
        out_html = render_html(doc_meta, blocks, page_meta, leg_nav)
        with open(os.path.join(LEG_DIR, f"{doc_meta['slug']}.html"), "w") as f:
            f.write(out_html)
        with open(os.path.join(OUT_DIR, f"{doc_meta['slug']}.html"), "w") as f:
            f.write(build_redirect(doc_meta["slug"]))
        counts = {t: sum(1 for b in blocks if b["type"] == t) for t in ("h2", "h3", "h4", "h5", "table")}
        print(doc_meta["slug"], "blocks=", len(blocks), counts, "meta=", page_meta)

    with open(os.path.join(LEG_DIR, "index.html"), "w") as f:
        f.write(build_legislation_index(leg_nav))
    with open(os.path.join(OUT_DIR, "index.html"), "w") as f:
        f.write(build_hub(manifest, hub_nav))
    with open(os.path.join(ASSETS_DIR, "legislation.css"), "w") as f:
        f.write(FONTS_CSS + theme_css + BASE_CSS + LEGISLATION_CSS)
    with open(os.path.join(ASSETS_DIR, "hub.css"), "w") as f:
        f.write(FONTS_CSS + theme_css + BASE_CSS + HUB_CSS)
    shutil.copyfile(
        os.path.join(KIT_DIR, "theme-toggle.js"), os.path.join(ASSETS_DIR, "theme-toggle.js")
    )
    with open(os.path.join(OUT_DIR, "sitemap.xml"), "w") as f:
        f.write(build_sitemap(manifest))

    print(f"\nhub, legislation index, {len(DOCS)} documents, {len(DOCS)} redirect stubs, assets, sitemap.")

    if unmapped_styles:
        print("\nWARNING: unmapped styles rendered as plain <p> (review CSS_CLASS_BY_STYLE):")
        for s in sorted(unmapped_styles):
            print(" -", s)
    else:
        print("All styles mapped.")


if __name__ == "__main__":
    main()
